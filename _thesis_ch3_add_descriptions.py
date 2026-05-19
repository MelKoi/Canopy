# -*- coding: utf-8 -*-
"""在第三章各功能需求小节「表3.x」前补充 1～2 段说明（基于《天幕》工程语境）。"""
import os
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _new_paragraph_xml(text: str):
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def _find_table_paragraph_index(doc: Document, caption: str) -> int:
    cap = caption.strip()
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == cap:
            return i
    raise KeyError(f"未找到段落：{caption!r}")


def _insert_two_paragraphs_before_table(doc: Document, table_caption: str, para1: str, para2: str):
    ti = _find_table_paragraph_index(doc, table_caption)
    if ti <= 0:
        raise RuntimeError(f"{table_caption} 前无段落")
    prev = doc.paragraphs[ti - 1].text.strip()
    if prev and prev.startswith(para1.strip()[:20]):
        print("skip (already)", table_caption)
        return
    doc.paragraphs[ti - 1].text = para1
    doc.paragraphs[ti]._element.addprevious(_new_paragraph_xml(para2))


def main():
    if len(sys.argv) < 2:
        print("usage: _thesis_ch3_add_descriptions.py <论文.docx>", file=sys.stderr)
        sys.exit(2)
    path = os.path.abspath(sys.argv[1])
    if not os.path.isfile(path):
        print("not found:", path, file=sys.stderr)
        sys.exit(1)

    blocks = [
        (
            "表3.1",
            "玩家移动是高速战斗的底座。《天幕》在工程上以机甲根刚体为核心，将水平推进、跳跃与空中姿态与核心能量、闪避无敌窗等参数绑定在同一套输入—状态机路径中，使「普通推进—闪避—极速推进」在手感上可区分、在资源上可制衡，并为后续关卡按 Prefab 调参留出接口。",
            "本需求表将位移速度、加速度、能量消耗与冷却等条目与上述模式一一对应：地面与空中分支在 MeshController 侧区分重力与贴地检测；闪避与极速推进在能量不足时自动降级或禁止，从而把策划案中的节奏目标落实为可测试的数值列。",
        ),
        (
            "表3.2",
            "武器射击需求对应玩家侧「四槽独立开火」与弹药经济。《天幕》中 WeaponRaycastShooter 以屏幕中心射线与 LockOnSystem 目标共同决定指向，按槽位解析 Hand/Shoulder 下装备的「炮口」「枪口」Transform；弹药则按装备根物体名称识别火箭筒、连发枪、单发枪/炮等档案，并支持换弹组合键与弹匣/备弹拆分。",
            "表内条目将射速、弹道可见性、火箭筒停步—发射硬直与对敌伤害系数等与上述实现挂钩：射线方案降低大量实体弹的 GC 与碰撞成本，同时保留尾迹与口径等可读参数，便于在需求评审中与性能指标对照。",
        ),
        (
            "表3.3",
            "战斗数值需求把「打得动、打得懂」落到生命、韧性与受击反馈上。工程里 PlayerMechResources 维护生命与韧性累积，敌人弹体命中时扣血并涨韧性，韧性满触发短时全身闪红后归零，生命为零销毁玩家；敌人侧 EnemyHitFeedback 提供可配置生命与延迟销毁，支撑关卡难度梯度。",
            "若论文前文将临界值—失衡作为循环核心，本表中的冲击、易伤窗口与保护期等字段即为其可量化拆解；实现上可与 UI 血条、受击音效同一事件链驱动，保证需求条目在联调时能被逐项勾选。",
        ),
        (
            "表3.4",
            "锁定需求减轻高速下持续精瞄的负担。《天幕》中 LockOnSystem 以准星视口半径筛选敌人实现软锁，并以鼠标中键进入硬锁、距离上限与屏幕外剔除做保护；武器与相机可读取 currentTarget 做轻微偏轴与辅助瞄准，避免完全自动锁死视角。",
            "表中距离、视口比例、硬锁拾取半径等数值与上述组件序列化字段一致：需求阶段用于约定「何时出现框、何时可开火偏轴」；测试阶段则与第七章手感条目对应，便于回归。",
        ),
        (
            "表3.5",
            "敌人 AI 需求强调在有限算力下仍能给高速玩家施压。工程以巡逻—交战挂起—冷却射击的组件组合为主：EnemyPatrolAgent 沿路径折线移动，TestEnemyCombat 等在接战条件成立时暂停巡逻并按间隔生成可见弹体；飞机 Boss 则由 PlaneDistanceHoverAI 维持环绕距离，PlaneCombat 分主副炮组驱动齐射与爆发冷却。",
            "表内感知半径、视线掩体、射速与伤害等列与上述脚本公开字段对齐，便于把「职能分工」留在数据层、把行为留在组件层；后续若引入统一 AI 表，可直接由本表映射到 ScriptableObject 行而不必重写交战管线。",
        ),
        (
            "表3.6",
            "非功能需求约束工程在真实硬件上的可交付性：除帧率、内存与加载时间外，还应考虑脚本热路径上的物理查询次数、弹体生成策略对 GC 的影响，以及本地化资源与中文排版在 UI 上的可读性。",
            "《天幕》作为 Unity 实时项目，将上述指标落实为「战斗主循环内避免每帧分配」「大场景分块加载与 LOD 预留」等可测条款，并与第七章性能测试数据闭环；本表为验收提供检查清单，而非替代具体Profiler结论。",
        ),
    ]

    for cap, a, b in blocks:
        doc = Document(path)
        _insert_two_paragraphs_before_table(doc, cap, a, b)
        doc.save(path)
        print("patched", cap)

    print("OK", path)


if __name__ == "__main__":
    main()
