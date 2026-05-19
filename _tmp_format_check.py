# -*- coding: utf-8 -*-
import re
import zipfile
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml.ns import qn

path = r"C:\Users\Narrator-鲤\Desktop\基于Unity3d的高速战斗系统制作.docx"
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
TW = 56.6929


def mm_twips(tw):
    return round(int(tw) / TW, 1)


def main():
    z = zipfile.ZipFile(path)
    root = ET.fromstring(z.read("word/document.xml"))
    body = root.find(W + "body")

    print("=== 4.1.1 页边距 / 页眉页脚距 (各节 sectPr) ===")
    for idx, sp in enumerate(body.iter(W + "sectPr")):
        pg = sp.find(W + "pgMar")
        pgt = sp.find(W + "pgNumType")
        if pg is not None:
            top_k = W + "top"
            bot_k = W + "bottom"
            left_k = W + "left"
            right_k = W + "right"
            hdr_k = W + "header"
            ftr_k = W + "footer"
            print(
                f"  节{idx}: 上{mm_twips(pg.get(top_k))} 下{mm_twips(pg.get(bot_k))} "
                f"左{mm_twips(pg.get(left_k))} 右{mm_twips(pg.get(right_k))} "
                f"页眉距{mm_twips(pg.get(hdr_k))} 页脚距{mm_twips(pg.get(ftr_k))} mm"
            )
        if pgt is not None:
            print("    pgNumType:", dict(pgt.attrib))

    print("\n=== 4.1.2 页眉 / 4.1.3 页脚页码 (header/footer xml) ===")
    want = "湖南科技大学本科生毕业设计（论文）"
    for fn in sorted(z.namelist()):
        if not fn.endswith(".xml") or "rels" in fn:
            continue
        if "header" not in fn and "footer" not in fn:
            continue
        troot = ET.fromstring(z.read(fn))
        texts = []
        instrs = []
        for el in troot.iter(W + "t"):
            if el.text:
                texts.append(el.text)
        for el in troot.iter(W + "instrText"):
            if el.text:
                instrs.append(el.text.strip())
        text = "".join(texts).replace("\n", "")
        if not text and not instrs:
            continue
        sz_pt = None
        for r in troot.iter(W + "r"):
            rpr = r.find(W + "rPr")
            if rpr is None:
                continue
            sz = rpr.find(W + "sz")
            rf = rpr.find(W + "rFonts")
            if sz is not None:
                sz_pt = int(sz.get(W + "val")) / 2
                ea = rf.get(W + "eastAsia") if rf is not None else None
                break
        jc = troot.find(".//" + W + "jc")
        align = jc.get(W + "val") if jc is not None else None
        print(f"  {fn}: text={text!r} fields={instrs} font_pt={sz_pt} align={align}")
        if want in text:
            if sz_pt and abs(sz_pt - 10.5) > 0.6:
                print("    !! 页眉字号非五号(10.5pt)")
            if align != "center":
                print("    !! 页眉未居中")

    doc = Document(path)
    breaks = []
    for i, p in enumerate(doc.paragraphs):
        p_pr = p._element.find(qn("w:pPr"))
        if p_pr is not None and p_pr.find(qn("w:sectPr")) is not None:
            breaks.append(i)
    print(f"\n分节符数量: {len(breaks)} (共 {len(breaks)+1} 节)")
    for pi in breaks:
        print(f"  段{pi}后分节: {repr((doc.paragraphs[pi].text or '')[:35])}")

    print("\n=== 4.2 样式库 (Heading/Normal) ===")
    expect = {"Heading 1": 18, "Heading 2": 14, "Heading 3": 12, "Normal": 12}
    for name, exp in expect.items():
        try:
            s = doc.styles[name]
            fs = s.font.size.pt if s.font.size else None
            pf = s.paragraph_format
            fi = pf.first_line_indent
            fi_pt = fi.pt if fi else 0
            ls = pf.line_spacing
            print(
                f"  {name}: {fs}pt (期望{exp}) bold={s.font.bold} "
                f"line_spacing={ls} first_indent={fi_pt}pt"
            )
        except KeyError:
            print(f"  {name}: 缺失")

    print("\n=== 4.9 摘要/关键词 抽样 ===")
    for i in [9, 10, 14, 16, 17, 22]:
        if i >= len(doc.paragraphs):
            continue
        p = doc.paragraphs[i]
        sizes = list({r.font.size.pt for r in p.runs if r.font.size})
        bolds = list({r.font.bold for r in p.runs if r.bold is not None})
        print(
            f"  段{i} {repr((p.text or '')[:22])} style={p.style.name} "
            f"sizes={sizes} bold={bolds} align={p.paragraph_format.alignment}"
        )

    print("\n=== 4.4 公式编号 ===")
    for i, p in enumerate(doc.paragraphs):
        t = p.text or ""
        if re.search(r"[\(（]\d+\.\d+[\)）]", t) and len(t) < 120:
            print(f"  段{i}: {t}")

    print("\n=== 4.5 表题 ===")
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if re.match(r"^表\s*[\d\.．]+", t):
            bad = any(x in t for x in "，。；：、")
            if "表 3.6 将" in t or bad or " " in t[:4]:
                print(f"  段{i}: {t[:55]}{' [非表题/含标点]' if bad or '将' in t else ''}")

    print("\n=== 4.6 图题 / 图号重复 ===")
    fig_nums = []
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        m = re.match(r"^(图\s*[\d\.]+)\s*(.*)", t)
        if m:
            fig_nums.append((m.group(1).replace(" ", ""), i, t))
    from collections import Counter

    c = Counter(x[0] for x in fig_nums)
    for num, cnt in c.items():
        if cnt > 1:
            locs = [x for x in fig_nums if x[0] == num]
            print(f"  重复 {num}: {cnt}次 -> {locs}")

    print("\n=== 4.7 参考文献 ===")
    p = doc.paragraphs[459]
    print(f"  标题段459: {repr(p.text)} style={p.style.name}")
    if p.runs and p.runs[0].font.size:
        print(f"    字号 {p.runs[0].font.size.pt}pt (期望标题四号14pt)")
    e = doc.paragraphs[460]
    print(f"  条目段460: size={e.runs[0].font.size.pt if e.runs and e.runs[0].font.size else None} (期望五号10.5pt)")
    print(f"    {e.text[:70]}")

    print("\n=== 4.3 章另起页 (pageBreakBefore) ===")
    for i in [97, 114, 159, 203, 230, 316, 407, 452]:
        p = doc.paragraphs[i]
        p_pr = p._element.find(qn("w:pPr"))
        pb = p_pr is not None and p_pr.find(qn("w:pageBreakBefore")) is not None
        print(f"  段{i} {p.text[:12]} pageBreakBefore={pb} style={p.style.name}")

    z.close()


if __name__ == "__main__":
    main()
