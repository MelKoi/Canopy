# -*- coding: utf-8 -*-
import zipfile
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml.ns import qn

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
path = r"C:\Users\Narrator-鲤\Downloads\基于Unity3d的高速战斗系统制作.docx"


def parse_part(data: bytes):
    root = ET.fromstring(data)
    texts, instrs, fonts = [], [], []
    for el in root.iter():
        if el.tag == W + "t" and el.text:
            texts.append(el.text)
        if el.tag == W + "instrText" and el.text:
            instrs.append(el.text.strip())
    for r in root.iter(W + "r"):
        rpr = r.find(W + "rPr")
        if rpr is not None:
            rf = rpr.find(W + "rFonts")
            sz = rpr.find(W + "sz")
            if rf is not None or sz is not None:
                fonts.append(
                    {
                        "ascii": rf.get(W + "ascii") if rf is not None else None,
                        "eastAsia": rf.get(W + "eastAsia") if rf is not None else None,
                        "sz_half_pt": sz.get(W + "val") if sz is not None else None,
                    }
                )
    jc = root.find(".//" + W + "jc")
    align = jc.get(W + "val") if jc is not None else None
    return "".join(texts), instrs, fonts[:4], align


z = zipfile.ZipFile(path)
print("=== HEADERS / FOOTERS ===")
for kind in ("header", "footer"):
    for i in range(1, 10):
        fn = f"word/{kind}{i}.xml"
        if fn not in z.namelist():
            continue
        t, ins, f, a = parse_part(z.read(fn))
        sz_pt = None
        if f and f[0].get("sz_half_pt"):
            sz_pt = int(f[0]["sz_half_pt"]) / 2
        print(f"{fn}: text={t!r} fields={ins} align={a} font_pt={sz_pt} fonts={f}")

doc = Document(path)
sect_indices = []
for i, p in enumerate(doc.paragraphs):
    p_pr = p._element.find(qn("w:pPr"))
    if p_pr is None:
        continue
    sp = p_pr.find(qn("w:sectPr"))
    if sp is not None:
        sect_indices.append(i)

print("\n=== SECTION BREAKS ===")
for si, pi in enumerate(sect_indices):
    print(f"Section break after para {pi}: {repr((doc.paragraphs[pi].text or '')[:40])}")
    sp = doc.paragraphs[pi]._element.find(qn("w:pPr")).find(qn("w:sectPr"))
    for ref in list(sp.findall(qn("w:headerReference"))) + list(
        sp.findall(qn("w:footerReference"))
    ):
        print(" ", ref.attrib)
    pg_num = sp.find(qn("w:pgNumType"))
    if pg_num is not None:
        print("  pgNumType:", pg_num.attrib)

# markers
print("\n=== CONTENT MARKERS ===")
for label, needle in [
    ("cover_end", "摘"),
    ("toc", "目"),
    ("ch1", "第一章"),
]:
    for i, p in enumerate(doc.paragraphs[:100]):
        t = (p.text or "").replace(" ", "")
        if needle in t and (label != "cover_end" or "要" in t):
            sec = 1
            for br in sect_indices:
                if i > br:
                    sec += 1
            print(f"  {label} para {i} section~{sec}: {repr((p.text or '')[:35])}")
            break

z.close()
