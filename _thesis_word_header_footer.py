# -*- coding: utf-8 -*-
"""
毕业设计 Word 页眉页脚与分节页码：
  第1节：仅封面 — 无页码；
  第2节：摘要～目录 — 小写罗马，样式「-i-」「-ii-」…，页脚 Times New Roman 五号居中；
  第3节：第一章起 — 阿拉伯数字连续「-1-」「-2-」…，同上。

依赖：Windows 已安装 Microsoft Word、pywin32、python-docx（仅用于移动分节符 XML）。
"""
from __future__ import annotations

import os
import sys
from copy import deepcopy

import win32com.client as win32
from docx import Document as DocxDocument
from docx.oxml.ns import qn

# Word 常量
WD_ALIGN_PARAGRAPH_CENTER = 1
WD_SECTION_BREAK_NEXT_PAGE = 2
WD_HEADER_FOOTER_PRIMARY = 1
WD_HEADER_FOOTER_FIRST_PAGE = 2
WD_COLLAPSE_START = 1
WD_FIND_STOP = 0

WD_PAGE_NUMBER_STYLE_LOWER_ROMAN = 2
WD_PAGE_NUMBER_STYLE_NUMBER_IN_DASH = 57

HEADER_TEXT = "湖南科技大学本科生毕业设计（论文）"
HEADER_FONT = "宋体"
HEADER_SIZE = 10.5

FOOTER_FONT = "Times New Roman"
FOOTER_SIZE = 10.5


def _clear_range(rng):
    rng.Text = ""
    rng.ParagraphFormat.Alignment = WD_ALIGN_PARAGRAPH_CENTER


def _apply_footer_font(rng):
    rng.Font.Name = FOOTER_FONT
    rng.Font.Size = FOOTER_SIZE
    try:
        rng.Font.NameAscii = FOOTER_FONT
    except Exception:
        pass


def _clear_footer_remove_fields(footer):
    rng = footer.Range
    try:
        for i in range(rng.Fields.Count, 0, -1):
            rng.Fields(i).Delete()
    except Exception:
        pass
    _clear_range(rng)


def relocate_first_section_break_after_cover(docx_path: str) -> bool:
    """
    将「目录前」的分节符移到「封面末、摘要前」（段落 12 末），使第 1 节仅为封面。
    若已是 [12,38] 两处分节则跳过。返回是否写回了文件。
    """
    doc = DocxDocument(docx_path)
    idx = []
    for i, p in enumerate(doc.paragraphs):
        p_pr = p._element.find(qn("w:pPr"))
        if p_pr is None:
            continue
        if p_pr.find(qn("w:sectPr")) is not None:
            idx.append(i)
    # 封面末、摘要前；第一章前（文档修订后索引可能为 [10,34] 或 [12,38]）
    if idx in ([12, 38], [10, 34], [8, 32], [8, 22], [8, 21]):
        return False
    # 仅封面末一处分节（摘要起为第 2 节），第一章前分节由页码脚本插入
    if len(idx) == 1 and idx[0] <= 20:
        return False
    if idx == [23, 38]:
        p23_pr = doc.paragraphs[23]._element.find(qn("w:pPr"))
        sect = p23_pr.find(qn("w:sectPr"))
        new_sect = deepcopy(sect)
        p23_pr.remove(sect)
        p12_pr = doc.paragraphs[12]._element.get_or_add_pPr()
        old = p12_pr.find(qn("w:sectPr"))
        if old is not None:
            p12_pr.remove(old)
        p12_pr.append(new_sect)
        doc.save(docx_path)
        return True
    if idx == [38]:
        return False
    raise RuntimeError(f"未识别的分节段落索引 {idx!r}，请手动检查分节符位置。")


def ensure_section_before_abstract(word_doc):
    """在「摘 要」前插入下一页分节符，使节数为 3（当仅有「第一章前」一处分节时）。"""
    if word_doc.Sections.Count >= 3:
        return
    find_rng = word_doc.Content
    f = find_rng.Find
    f.ClearFormatting()
    f.Text = "摘 要"
    f.Forward = True
    f.Wrap = WD_FIND_STOP
    if not f.Execute():
        raise RuntimeError("未找到「摘 要」标题，无法插入分节符。")
    rng = find_rng.Paragraphs(1).Range
    rng.Collapse(WD_COLLAPSE_START)
    rng.InsertBreak(WD_SECTION_BREAK_NEXT_PAGE)


def set_footer_roman_dash(sect):
    """第2节：罗马数字，节内从 1 起；两侧 ASCII 短横线；五号 TNR 居中。"""
    ft = sect.Footers(WD_HEADER_FOOTER_PRIMARY)
    ft.LinkToPrevious = False
    _clear_footer_remove_fields(ft)

    r = ft.Range
    r.ParagraphFormat.Alignment = WD_ALIGN_PARAGRAPH_CENTER
    pns = ft.PageNumbers
    try:
        pns.RestartNumberingAtSection = True
    except Exception:
        pass
    try:
        pns.StartingNumber = 1
    except Exception:
        pass
    try:
        pns.Add(PageNumberAlignment=WD_ALIGN_PARAGRAPH_CENTER)
    except Exception:
        pass
    try:
        pns.NumberStyle = WD_PAGE_NUMBER_STYLE_LOWER_ROMAN
    except Exception:
        pass
    _apply_footer_font(ft.Range)
    try:
        r = ft.Range
        t = (r.Text or "").strip()
        if t and not t.startswith("-"):
            r.InsertBefore("-")
        if t and not t.endswith("-"):
            r.InsertAfter("-")
        _apply_footer_font(ft.Range)
    except Exception:
        pass
    for fld in ft.Range.Fields:
        try:
            fld.Update()
        except Exception:
            pass


def set_footer_arabic_dash(sect):
    """第3节：阿拉伯数字，-1- 样式，节内从 1 起。"""
    ft = sect.Footers(WD_HEADER_FOOTER_PRIMARY)
    ft.LinkToPrevious = False
    _clear_footer_remove_fields(ft)

    r = ft.Range
    r.ParagraphFormat.Alignment = WD_ALIGN_PARAGRAPH_CENTER
    pns = ft.PageNumbers
    try:
        pns.RestartNumberingAtSection = True
    except Exception:
        pass
    try:
        pns.StartingNumber = 1
    except Exception:
        pass
    try:
        pns.Add(PageNumberAlignment=WD_ALIGN_PARAGRAPH_CENTER)
    except Exception:
        pass
    try:
        pns.NumberStyle = WD_PAGE_NUMBER_STYLE_NUMBER_IN_DASH
    except Exception:
        pns.NumberStyle = 0
    _apply_footer_font(ft.Range)
    for fld in ft.Range.Fields:
        try:
            fld.Update()
        except Exception:
            pass


def configure_section1_cover_only(sec1):
    """第1节：封面，无页码、无页眉。"""
    try:
        sec1.PageSetup.DifferentFirstPageHeaderFooter = False
    except Exception:
        pass
    for hf in (WD_HEADER_FOOTER_PRIMARY, WD_HEADER_FOOTER_FIRST_PAGE):
        try:
            sec1.Headers(hf).LinkToPrevious = False
        except Exception:
            pass
        _clear_range(sec1.Headers(hf).Range)
    for ft_idx in (WD_HEADER_FOOTER_PRIMARY, WD_HEADER_FOOTER_FIRST_PAGE):
        try:
            sec1.Footers(ft_idx).LinkToPrevious = False
        except Exception:
            pass
        try:
            _clear_footer_remove_fields(sec1.Footers(ft_idx))
        except Exception:
            pass


def configure_section2_abstract_through_toc(sec2):
    """第2节：摘要～目录；页眉学校全称；罗马页码。"""
    try:
        sec2.PageSetup.DifferentFirstPageHeaderFooter = False
    except Exception:
        pass

    for hf in (WD_HEADER_FOOTER_PRIMARY, WD_HEADER_FOOTER_FIRST_PAGE):
        try:
            sec2.Headers(hf).LinkToPrevious = False
        except Exception:
            pass
        _clear_range(sec2.Headers(hf).Range)

    hr = sec2.Headers(WD_HEADER_FOOTER_PRIMARY).Range
    hr.Text = HEADER_TEXT
    hr.ParagraphFormat.Alignment = WD_ALIGN_PARAGRAPH_CENTER
    hr.Font.Name = HEADER_FONT
    hr.Font.Size = HEADER_SIZE
    try:
        hr.Font.NameFarEast = HEADER_FONT
    except Exception:
        pass

    try:
        sec2.Footers(WD_HEADER_FOOTER_FIRST_PAGE).LinkToPrevious = False
        _clear_footer_remove_fields(sec2.Footers(WD_HEADER_FOOTER_FIRST_PAGE))
    except Exception:
        pass

    set_footer_roman_dash(sec2)


def configure_section3_body(sec3):
    """第3节：正文；页眉接续第2节；阿拉伯 -1- 起编。"""
    try:
        sec3.PageSetup.DifferentFirstPageHeaderFooter = False
    except Exception:
        pass
    try:
        sec3.Headers(WD_HEADER_FOOTER_PRIMARY).LinkToPrevious = True
    except Exception:
        pass
    try:
        h2 = sec3.Headers(WD_HEADER_FOOTER_FIRST_PAGE)
        h2.LinkToPrevious = False
        _clear_range(h2.Range)
    except Exception:
        pass
    set_footer_arabic_dash(sec3)


def main():
    if len(sys.argv) < 2:
        print("usage: _thesis_word_header_footer.py <论文.docx>", file=sys.stderr)
        sys.exit(2)
    path = os.path.abspath(sys.argv[1])
    if not os.path.isfile(path):
        print("file not found:", path, file=sys.stderr)
        sys.exit(1)

    # 先调整分节符 XML（避免 Word 独占时失败）
    relocate_first_section_break_after_cover(path)

    app = win32.Dispatch("Word.Application")
    app.Visible = False
    try:
        app.DisplayAlerts = 0
    except Exception:
        pass

    doc = None
    try:
        doc = app.Documents.Open(path, ReadOnly=False, AddToRecentFiles=False)
        ensure_section_before_abstract(doc)
        if doc.Sections.Count < 3:
            raise RuntimeError(
                f"分节后节数为 {doc.Sections.Count}，需要 3 节（封面 / 摘要～目录 / 正文）。"
            )

        configure_section1_cover_only(doc.Sections(1))
        configure_section2_abstract_through_toc(doc.Sections(2))
        configure_section3_body(doc.Sections(3))

        try:
            doc.Fields.Update()
        except Exception:
            pass
        doc.Save()
        print("OK", path)
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        try:
            app.Quit()
        except Exception:
            pass


if __name__ == "__main__":
    main()
