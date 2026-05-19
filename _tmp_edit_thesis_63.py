# -*- coding: utf-8 -*-
"""One-off: rewrite thesis section 6.3 in-place (思路+落地期望+结果后续)."""
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def paragraph_text(p_el):
    parts = []
    for t in p_el.iter(qn("w:t")):
        if t.text:
            parts.append(t.text)
    return "".join(parts)


def new_paragraph():
    return OxmlElement("w:p")


def add_text_run(p_el, text, bold=False):
    r = OxmlElement("w:r")
    if bold:
        rpr = OxmlElement("w:rPr")
        b = OxmlElement("w:b")
        rpr.append(b)
        r.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p_el.append(r)


def add_heading_line(p_el, title):
    """Single run bold line."""
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rpr.append(OxmlElement("w:b"))
    r.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = title
    r.append(t)
    p_el.append(r)


def add_triplet_block(body, insert_before_el, heading, idea, expect, follow):
    """Insert 4 paragraphs before 6.4: heading, 思路, 落地期望, 结果后续 (top to bottom)."""
    # Each insert(..., insert_before_el) puts new node immediately before anchor;
    # inserting in order 结果→落地→思路→title yields title on top, 结果 closest to 6.4.
    blocks = [
        ("结果后续：" + follow, False),
        ("落地期望：" + expect, False),
        ("思路：" + idea, False),
        (heading, True),
    ]
    for title, is_title in blocks:
        p = new_paragraph()
        if is_title:
            add_heading_line(p, title)
        else:
            add_text_run(p, title)
        idx = list(body).index(insert_before_el)
        body.insert(idx, p)


def main():
    if len(sys.argv) < 2:
        print("usage: script.py <input.docx> [output.docx]", file=sys.stderr)
        sys.exit(2)
    path = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else path
    doc = Document(path)
    body = doc.element.body
    children = list(body)

    def is_standalone_heading(tx, heading_start, must_contain):
        """避免匹配目录长段：仅匹配单独成段的章节标题。"""
        if not tx:
            return False
        t = tx.strip().replace("\u00a0", " ")
        if "\n" in t:
            return False
        if len(t) > 120:
            return False
        if not t.startswith(heading_start + " "):
            return False
        return must_contain in t

    idx_63_title = None
    idx_64 = None
    for i, el in enumerate(children):
        if el.tag != qn("w:p"):
            continue
        tx = paragraph_text(el)
        if idx_63_title is None and is_standalone_heading(tx, "6.3", "敌人攻击系统设计"):
            idx_63_title = i
            continue
        if idx_63_title is not None and is_standalone_heading(tx, "6.4", "敌人行为状态机设计"):
            idx_64 = i
            break

    if idx_63_title is None or idx_64 is None or idx_64 <= idx_63_title:
        print("Could not find standalone 6.3 / 6.4 headings", idx_63_title, idx_64, file=sys.stderr)
        sys.exit(1)

    insert_point = children[idx_64]
    # Remove everything after 6.3 title until (not including) 6.4
    for j in range(idx_64 - 1, idx_63_title, -1):
        body.remove(children[j])

    # Insert new content (order: insert at same index repeatedly pushes 6.4 down)
    # We insert_before insert_point; each insert shifts index — insert from bottom content first
    sections = [
        (
            "6.3.1 弹道预测",
            "高速机甲战中玩家位移大，若始终瞄当前世界坐标，弹体飞行时间内目标已离开命中域；因此需要在发射前用目标速度做一阶或迭代前置量，并可混入精度系数以控制难度曲线。",
            "在《天幕》工程里，将玩家水平速度由 Rigidbody 或 MechController 读出，在敌人开火前计算 travelTime≈距离/弹速，并令瞄准点沿速度方向外推；对飞机 Boss 与小兵可配置不同预测强度与随机抖动表，以便关卡单独调参。",
            "当前仓库中 TestEnemyCombat、PlaneCombat 等仍以当前时刻指向为主，前置量尚未作为默认路径接入；按上式补齐后，可在不改弹体Prefab的前提下显著改善高速目标命中率，并为后续「不同敌种预测系数」表提供数据回填入口。",
        ),
        (
            "6.3.2 AI 武器配置",
            "不同职能敌人应在射程、射速、弹种与停火窗口上拉开差异，使玩家形成距离感与资源交换；配置应尽量数据化，避免把数值写死在脚本常量里。",
            "结合现有实现：地面测试敌人以序列化字段控制检测半径、开火间隔、弹速与弹径；飞机 Boss 由 PlaneCombat 分 GUN 管与冷却、射程类参数；玩家侧 WeaponRaycastShooter 已按装备根物体名称区分单发/连发/火箭筒弹药与伤害系数——敌人侧可复用同一套「名称/标签驱动」或 ScriptableObject 行表，与表6.1 中轻型/重型/支援的职能划分一一对应。",
            "论文中原表6.1 的数值可保留为设计参照，正文以工程组件字段说明替代纯虚构表项；落地后可在同一关卡内并置多预制体实例，用 Inspector 差异体现三档职能，测试记录再回写论文数据列。",
        ),
        (
            "6.3.3 攻击实现",
            "攻击管线应短：交战成立 → 取瞄准点/枪口 → 生成弹体或射线 → 命中层过滤己方碰撞体 → 调用玩家受伤接口 → 进入冷却；火箭等重武器可插入短移动锁避免滑步开火。",
            "工程已具备关键链路：TestEnemyCombat 在冷却结束时生成可见弹体并挂 EnemyProjectileBullet，命中链路到 PlayerMechResources.RegisterEnemyProjectileHit；PlaneCombat 按接战距离驱动多组枪口齐射；WeaponRaycastShooter 中火箭筒协程展示「停步—发射—短硬直」模板，可平移到敌人重火力。",
            "后续若统一敌人武器表，可将「是否射线/是否实体弹」抽象为同一发射器接口，并把命中回调与 UI 受击反馈、音效触发绑在同一事件上，便于第七章功能测试用例逐条勾选。",
        ),
    ]

    # Insert in reverse order so final doc reads 6.3.1, 6.3.2, 6.3.3
    for heading, idea, expect, follow in reversed(sections):
        add_triplet_block(body, insert_point, heading, idea, expect, follow)

    doc.save(out)
    print("OK saved", out)


if __name__ == "__main__":
    main()
0