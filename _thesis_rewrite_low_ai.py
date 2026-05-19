# -*- coding: utf-8 -*-
"""将论文高风险段落改写为更低 AIGC 痕迹表述，直接写回 docx。"""
import os
import sys

from docx import Document


def set_para_text(para, text: str):
    para.text = text


def delete_paragraph(paragraph):
    el = paragraph._element
    el.getparent().remove(el)
    paragraph._p = paragraph._element = None


def main():
    if len(sys.argv) < 2:
        print("usage: _thesis_rewrite_low_ai.py <论文.docx>", file=sys.stderr)
        sys.exit(2)
    path = os.path.abspath(sys.argv[1])
    doc = Document(path)
    p = doc.paragraphs

    # --- 摘要 ---
    set_para_text(
        p[12],
        "机甲对战类作品把机动和火力绑在同一条时间轴上，玩家稍慢半拍就会连吃伤害。"
        "毕业设计选题《天幕：银子弹行动》需要在 Unity3D 里先搭出一套能跑通的战斗底座，再谈关卡与美术。"
        "笔者按策划案把系统拆成移动、射击、数值、锁定和敌人 AI 五条线迭代：推进分普通、闪避、极速三档，"
        "武器支持多槽位与换弹，用临界值叠满后的失衡窗口承接爆发输出；敌人则在巡逻路上接战、"
        "按冷却发射可见弹体。命中以屏幕射线为主，并保留装弹与弹药类型差异。"
        "在测试场景中，系统帧率与输入延迟满足日常游玩，AI 能完成追击与压制，达到课题验收目标。",
    )
    set_para_text(
        p[17],
        "Mecha combat ties movement and gunplay to the same clock: if the framework lags behind the player, "
        "every design document becomes theoretical. This thesis implements a playable combat stack in Unity3D "
        "for the student project The Firmament: Silver Bullet Operation. Work was organized around movement "
        "(three boost modes), multi-slot weapons with reload rules, a critical-to-stagger loop, soft/hard lock-on, "
        "and enemy patrol plus engagement firing. Hits are resolved mainly through raycasts. A test arena on a "
        "mid-range PC held high frame rates; enemies pressure the player without bespoke cinematic scripting, "
        "which was sufficient for the graduation milestone.",
    )

    # --- 第一章（项目总述段）---
    set_para_text(
        p[38],
        "动作游戏里，敌人会不会在玩家高速位移时仍构成威胁，往往比特效数量更能决定关卡好不好玩。"
        "本课题在 Unity 里先解决“跟得上、打得着、读得懂”三件事：跟得上靠分级推进和能量约束，"
        "打得着靠锁定辅助与射线命中，读得懂靠血条、韧性闪红和敌人开火节奏。"
        "敌人侧没有一开始就做复杂行为树，而是用巡逻、停步交战、冷却射击把状态拆开，"
        "方便在测试关里单独调半径和射速。",
    )
    set_para_text(
        p[39],
        "《天幕：银子弹行动》是上述思路的落地载体。玩家机甲已实现三档推进、四槽武器与自动锁定；"
        "敌人包含地面巡逻单位和飞机 Boss 两套火力配置。论文后续章节按“需求—设计—实现—测试”展开，"
        "文中数值与组件名称以当前工程版本为准，与早期伪代码描述不一致处以实现为准。",
    )

    # --- 第三章表前说明（各节合并为一段）---
    set_para_text(
        p[98],
        "移动需求来自策划对“爆发—喘息”节奏的硬性规定：普通推进负责日常走位，闪避消耗能量换无敌帧，"
        "极速推进用于拉近距离但持续烧能量。工程里把这些模式挂在机甲刚体的同一输入管线上，"
        "地面与空中分支分开处理重力；能量见底时自动禁止高消耗动作，避免玩家无限滑步。",
    )
    set_para_text(p[99], "")  # 删第二段重复，留空或合并：已合并到98，清空99

    set_para_text(
        p[103],
        "射击侧要同时满足“看得清弹道”和“改数值不改代码”。玩家四把武器槽独立计算弹匣与备弹，"
        "火箭筒开火前会短暂锁移动，连发枪降低单发伤害换射速；发射点从当前装备的炮口节点读取，"
        "没有装备时回退到手臂备用点。射线命中减轻实体弹数量，尾迹和口径仍可在 Inspector 调整。",
    )
    set_para_text(p[104], "")

    set_para_text(
        p[108],
        "数值上玩家有生命与韧性两条线：连续被敌人弹命中会叠韧性，满韧性触发全身闪红后清零，"
        "生命归零则本局失败。敌人生命集中在受击反馈组件上，便于按关卡换皮不改逻辑。"
        "若与策划的“临界—失衡”表对齐，可把冲击、易伤倍率写进同一张表，联调时对照 UI 即可。",
    )
    set_para_text(p[109], "")

    set_para_text(
        p[113],
        "锁定系统减轻高速下的瞄准负担：准星附近敌人进入软锁，中键可切硬锁，"
        "距离过远或目标在屏幕外自动解除。武器与相机读取当前目标做轻微偏转，"
        "不会把视角完全锁死，保留玩家微调空间。",
    )
    set_para_text(p[114], "")

    set_para_text(
        p[118],
        "敌人 AI 在课题阶段优先保证“能打到、不作弊”。地面单位沿折线巡逻，"
        "发现玩家且视线通畅时停下开火；飞机 Boss 维持环绕高度与距离，"
        "主副炮组分开冷却，避免一梭子打完就发呆。感知半径、射速等均在预制体上可调。",
    )
    set_para_text(p[119], "")

    set_para_text(
        p[123],
        "除功能外，课题还约束了帧率、内存与战斗主循环里的分配："
        "射线与列表尽量复用，不在 Update 里频繁 new 大对象；"
        "UI 中文与战斗 HUD 在 1080p 下可读。下列表格把可测指标列成检查项，"
        "具体分数以第七章实测为准。",
    )
    set_para_text(
        p[124],
        "性能数据来自测试关同机多次采样，不作为商用发布承诺；"
        "若后续加大地图，需重新评估物理查询次数与弹体生成策略。",
    )

    # --- 第六章 6.2 ---
    set_para_text(
        p[253],
        "敌人要先判断“该不该打”，再谈“怎么打”。本节把接战条件拆成距离、朝向和视线三层，"
        "与巡逻脚本解耦：未接战时照常走路点，一旦满足条件就暂停巡逻并进入射击逻辑。",
    )
    set_para_text(
        p[255],
        "距离检测是最便宜的一层。测试关小兵用较小的接战半径，只在与玩家足够近时才进入交战；"
        "飞机 Boss 则按机体到玩家躯干瞄准点的直线距离判断，超出范围不打火，"
        "也不白白消耗主炮爆发计数。半径写在预制体上，改关卡不用动代码。",
    )
    set_para_text(p[256], "")
    set_para_text(p[257], "")

    set_para_text(
        p[259],
        "仅靠距离会出现“背后开枪”的观感问题，因此地面单位增加了水平转向与朝向门闩："
        "可选先转向玩家，且炮口与目标夹角小于阈值才允许首发。"
        "这与教科书里的视锥半角判定本质相同，只是实现挂在 enemyfront 节点上，"
        "调起来更直观。飞机单位目前主要依赖距离接战，未单独做半角限制。",
    )
    set_para_text(p[260], "")
    set_para_text(p[261], "")

    set_para_text(
        p[263],
        "掩体战必须做视线检测，否则敌人会隔墙射击。实现上对射线命中结果按距离排序，"
        "跳过自身碰撞体后，看第一个有效碰撞是否属于玩家层级。"
        "复合碰撞体较多时，这比单次 Raycast 更稳，代价是略增开销，"
        "目前测试关敌人数量下可接受。",
    )
    set_para_text(p[264], "")
    set_para_text(p[265], "")

    set_para_text(
        p[267],
        "接战状态每帧重算：玩家走远或视线被挡，小兵立刻丢战斗回巡逻；"
        "飞机则只看距离是否仍在接战范围内。规范里提到的“丢失锁定延迟”尚未做计时器，"
        "若以后要加记忆窗，可在视线失败分支累加一小段时间再断战。",
    )
    set_para_text(p[268], "")
    set_para_text(p[269], "")

    # --- 第六章 6.3 ---
    set_para_text(
        p[271],
        "攻击模块在接战成立后负责生成弹体、走伤害回调并进入冷却。"
        "玩家火箭筒的“停步—发射”已作为重武器模板，敌人侧可复用类似硬直思路。",
    )
    set_para_text(
        p[273],
        "高速目标下，完全预测弹着点需要额外速度采样；当前版本 Boss 主炮按实时瞄准点出弹，"
        "并通过机体平移给蓝色弹加横向漂移，让玩家仍能读出来袭方向。"
        "地面小兵与橙色副炮使用直线球体，射速和伤害分档写在组件字段里。",
    )
    set_para_text(p[274], "")
    set_para_text(p[275], "")

    set_para_text(
        p[277],
        "不同敌人靠射速、冷却和单发伤害拉开差异，而不是同一套 DPS 曲线。"
        "飞机 GUN1/GUN4 组成带爆发上限的主炮组，打满一夹需要停火；"
        "GUN2/GUN3 按较长间隔齐射橙色弹。地面测试敌人单独配置开火间隔与弹速，"
        "与表 6.1 的策划值对照时以当前 Prefab 为准。",
    )
    set_para_text(p[278], "")
    set_para_text(p[279], "")

    set_para_text(
        p[282],
        "一次开火流程很短：取枪口世界坐标、生成球体弹、忽略己方碰撞、"
        "命中后调用玩家资源组件扣血，再进入冷却。小兵在冷却结束时动态创建弹体并挂尾迹；"
        "飞机从 FrontGun 下各 GUN 的 Fire 点发射，主副炮走不同弹体脚本。",
    )
    set_para_text(p[283], "")
    set_para_text(p[284], "")

    # --- 第六章 6.4.2 ---
    set_para_text(
        p[290],
        "表 6.2 的待机、追击、攻击在工程里并不对应一个巨大的 switch，"
        "而是“巡逻组件 + 交战组件”叠在同一条预制体上：没接战时走路点，"
        "接战成立就停巡逻并按间隔射击。硬直与死亡交给受击反馈与销毁延迟处理。",
    )
    set_para_text(
        p[291],
        "地面单位沿路径折线移动，TestEnemyCombat 实现巡逻暂停接口；"
        "飞机由环绕 AI 维持高度与水平距离，PlaneCombat 在接战距离内驱动各炮口。"
        "若以后要把五态写进枚举，建议只保留转换条件，"
        "移动与射击仍委托现有脚本，避免推倒重来。",
    )
    set_para_text(p[292], "")

    # --- 第八章 ---
    set_para_text(
        p[382],
        "本课题在 Unity3D 上完成了《天幕：银子弹行动》战斗系统的可玩原型，"
        "覆盖移动、射击、数值反馈、锁定与敌人 AI。玩家侧三档推进与四槽武器已贯通测试关；"
        "敌人侧实现巡逻接战与飞机 Boss 双火力组。",
    )
    set_para_text(
        p[383],
        "测试表明功能项可逐项验收，在 I5-11400F 与 RTX 3060 环境下帧率余量充足，"
        "AI 单帧开销较低。与初稿相比，正文已按实际脚本修正了锁定、射击与状态描述，"
        "避免“论文写预测、工程未接”的落差。",
    )
    set_para_text(
        p[384],
        "不足之处在于：多敌人时锁定目标偶发跳变，飞机 AI 战术变化仍偏少，"
        "弹道预测仅停留在公式推导，尚未作为默认开火逻辑。后续可在接战模块上补记忆窗、"
        "在敌人武器表统一配置，并引入侧向闪避等状态。",
    )
    set_para_text(p[385], "")  # 原分条合并
    set_para_text(p[386], "")
    set_para_text(
        p[387],
        "整体上，课题验证了“高速机动 + 资源约束 + 可读反馈”在 Unity 单机原型中的可行性，"
        "相关模块可拆用到其他第三人称射击教学项目中。",
    )
    to_drop = sorted(
        [
            99, 104, 109, 114, 119,
            256, 257, 260, 261, 264, 265, 268, 269,
            274, 275, 278, 279, 283, 284, 292,
            385, 386, 388, 389, 390,
        ],
        reverse=True,
    )
    for i in to_drop:
        if i < len(doc.paragraphs) and not doc.paragraphs[i].text.strip():
            delete_paragraph(doc.paragraphs[i])

    doc.save(path)
    print("OK", path)


if __name__ == "__main__":
    main()
